from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def create_portfolio_document():
    # Créer un nouveau document
    doc = Document()

    # Définir les styles personnalisés
    styles = doc.styles

    # Titre principal
    title = doc.add_heading('Projet TP BIHAR 2025', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sous-titre
    subtitle = doc.add_paragraph()
    subtitle.add_run('Portfolio - Double Master BIHAR & Ingénierie Logicielle\n').bold = True
    subtitle.add_run('Vincent CHABRAN').italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Section Vue d'ensemble
    doc.add_heading('🎯 Vue d\'ensemble', 1)
    doc.add_paragraph(
        'Ce projet constitue un projet complet de Machine Learning et MLOps développé dans le cadre '
        'du double master BIHAR et Ingénierie Logicielle. Il démontre la maîtrise de trois domaines '
        'fondamentaux du ML : la vision par ordinateur, le traitement du langage naturel et l\'analyse '
        'de séries temporelles, avec une implémentation complète d\'une chaîne MLOps.'
    )
    doc.add_paragraph()

    # Section Compétences techniques
    doc.add_heading('🛠️ Compétences techniques démontrées', 1)

    # 1. Deep Learning & Computer Vision
    doc.add_heading('1. Deep Learning & Computer Vision', 2)
    cv_items = [
        'Architectures CNN : Implémentation d\'un CNN baseline from scratch et utilisation de transfer learning avec VGG16',
        'Prétraitement d\'images : Augmentation de données, normalisation, transforms PyTorch',
        'Classification multi-classes : Classification d\'images agricoles (maïs, herbes, sol) avec gestion de 3 et 4 classes',
        'Évaluation : Matrices de confusion, courbes d\'apprentissage, analyse des métriques (accuracy jusqu\'à 86% avec VGG16)'
    ]
    for item in cv_items:
        doc.add_paragraph(item, style='List Bullet')

    # 2. Natural Language Processing
    doc.add_heading('2. Natural Language Processing', 2)
    nlp_items = [
        'Vectorisation de texte : TF-IDF, Word2Vec (modèles pré-entraînés français)',
        'Modèles classiques : Régression logistique, SVM linéaire (92% d\'accuracy)',
        'Deep Learning NLP : Architecture LSTM pour l\'analyse de sentiments',
        'Prétraitement NLP : Nettoyage de texte, suppression des stopwords, tokenization',
        'Dataset : Travail sur le dataset Allocine (200k critiques de films en français)'
    ]
    for item in nlp_items:
        doc.add_paragraph(item, style='List Bullet')

    # 3. Séries temporelles
    doc.add_heading('3. Séries temporelles & Prévision', 2)
    ts_items = [
        'Modèles statistiques : ARIMA, SARIMA, SARIMAX avec recherche automatique des hyperparamètres',
        'Machine Learning : Random Forest (R² = 0.97), Régression linéaire',
        'Feature engineering : Création de features temporelles (lags, moyennes mobiles, variables calendaires)',
        'Analyse statistique : Tests de stationnarité (ADF, KPSS), ACF/PACF, décomposition STL',
        'Données météo : Intégration API Open-Meteo, gestion des variables exogènes'
    ]
    for item in ts_items:
        doc.add_paragraph(item, style='List Bullet')

    # 4. MLOps & Infrastructure
    doc.add_heading('4. MLOps & Infrastructure', 2)
    mlops_items = [
        'API REST : FastAPI avec endpoints documentés (Swagger)',
        'Base de données : SQLite avec ORM SQLAlchemy pour stockage des modèles et prédictions',
        'Monitoring : Scripts de comparaison prédictions/observations, génération de graphiques',
        'CI/CD : Pipeline GitHub Actions pour tests automatisés et déploiement',
        'Versioning : Gestion des versions de modèles dans la DB',
        'Logging : Système de logs structurés pour l\'API'
    ]
    for item in mlops_items:
        doc.add_paragraph(item, style='List Bullet')

    # 5. Engineering Best Practices
    doc.add_heading('5. Engineering Best Practices', 2)
    eng_items = [
        'Architecture modulaire : Séparation claire src/, notebooks/, data/, api/, monitoring/',
        'Tests unitaires : pytest pour validation de l\'API',
        'Configuration : Fichiers YAML pour paramètres des modèles',
        'Documentation : README détaillé, notebooks commentés, docstrings',
        'Containerisation : Support Docker pour déploiement'
    ]
    for item in eng_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # Section Résultats
    doc.add_heading('📈 Résultats notables', 1)

    # Créer un tableau des résultats
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light List Accent 1'

    # En-têtes du tableau
    headers = ['Domaine', 'Modèle', 'Performance']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    # Données du tableau
    data = [
        ['Vision', 'VGG16 (3 classes)', '86% accuracy'],
        ['NLP', 'TF-IDF + LogReg', '92% accuracy'],
        ['Séries temporelles', 'Random Forest', 'R² = 0.97, RMSE = 0.91'],
        ['API', 'FastAPI', '3 endpoints REST fonctionnels']
    ]

    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_paragraph()

    # Points d'innovation
    doc.add_heading('🔬 Points d\'innovation', 1)
    innovations = [
        'Pipeline MLOps complet : De l\'acquisition des données au déploiement API',
        'Comparaison multi-modèles : Benchmark systématique (5 modèles pour les séries temporelles)',
        'Gestion de la stationnarité : Transformation automatique et tests statistiques',
        'Feature engineering avancé : Variables exogènes météo + features temporelles',
        'Architecture évolutive : Base de données pour tracking des expériences'
    ]

    for i, innovation in enumerate(innovations, start=1):
        doc.add_paragraph(f'{i}. {innovation}')

    # Stack technologique
    doc.add_heading('📚 Stack technologique', 1)
    stack_items = {
        'Python': 'pandas, numpy, scikit-learn, statsmodels',
        'Deep Learning': 'PyTorch, TensorFlow/Keras',
        'MLOps': 'FastAPI, SQLAlchemy, Docker, GitHub Actions',
        'Visualisation': 'matplotlib, seaborn',
        'NLP': 'NLTK, gensim, transformers',
        'Tests': 'pytest, httpx'
    }

    for category, tools in stack_items.items():
        p = doc.add_paragraph()
        p.add_run(f'{category}: ').bold = True
        p.add_run(tools)

    doc.add_paragraph()

    # Compétences validées
    doc.add_heading('🎓 Compétences validées pour le master', 1)
    skills = [
        'Machine Learning avancé : Maîtrise des algorithmes classiques et deep learning',
        'MLOps : Déploiement et monitoring de modèles en production',
        'Analyse de données : EDA, feature engineering, validation croisée',
        'Développement logiciel : Clean code, tests, CI/CD',
        'Gestion de projet : Documentation, versioning Git, architecture modulaire'
    ]

    for skill in skills:
        p = doc.add_paragraph()
        p.add_run('✅ ').bold = True
        parts = skill.split(' : ')
        p.add_run(parts[0] + ' : ').bold = True
        if len(parts) > 1:
            p.add_run(parts[1])

    doc.add_page_break()

    # Architecture du projet
    doc.add_heading('📁 Architecture du projet', 1)
    doc.add_paragraph('Structure complète du repository :')

    architecture = """
notebooks/
    ├── image_classification.ipynb      # CNNs sur des images
    ├── text_classification.ipynb       # LSTM / analyse de sentiments
    └── timeseries_forecasting.ipynb    # Prévision de température

src/
    ├── image_classification/           # Modules pour la vision
    ├── text_classification/            # Modules pour le NLP
    ├── timeseries_forecasting/         # Modules pour les séries temporelles
    └── ForecastDatabase.py             # Interface base de données

data/
    ├── forecast_results.db             # Base SQLite
    └── configs/                        # Configurations YAML des modèles

api/
    ├── main.py                         # API FastAPI
    └── logs/                           # Journaux de l'API

monitoring/
    ├── compare_predictions.py          # Comparaison prédictions/observations
    └── output/                         # Graphiques générés

tests/
    ├── test_predict.py                 # Tests unitaires
    └── test_combined.py                # Tests d'intégration
    """

    doc.add_paragraph(architecture)

    # Lien GitHub
    doc.add_heading('🔗 Liens et ressources', 1)
    p = doc.add_paragraph()
    p.add_run('Repository GitHub : ').bold = True
    p.add_run('https://github.com/vincentchabran/tp-bihar-2025')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Date de génération : ').italic = True
    p.add_run(datetime.now().strftime('%d/%m/%Y'))

    # Sauvegarder le document
    filename = 'Portfolio_TP_BIHAR_2025.docx'
    doc.save(filename)
    print(f"✅ Document créé avec succès : {filename}")
    return filename

if __name__ == "__main__":
    create_portfolio_document()